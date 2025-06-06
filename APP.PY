# -*- coding: utf-8 -*-
"""
Flask Agenda Alignment Analyzer
Created on May 07, 2025
@author: Adity
"""

from flask import Flask, render_template, request, send_file, redirect, url_for, session
import moviepy.editor as mp
import whisper
import os
import re
import numpy as np
from sklearn.cluster import KMeans
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import RGBColor
from werkzeug.utils import secure_filename
import tempfile
import uuid
from datetime import datetime

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'Uploads'
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB limit
app.secret_key = 'your-secret-key-here'  # Change this in production

# Load models globally
model_whisper = whisper.load_model("medium")
model_st = SentenceTransformer('all-mpnet-base-v2')

# Store temporary report paths for download
report_paths = {}
# Store analysis results
analysis_results = {}

# Helper Functions
def extract_audio_from_video(video_path, audio_path):
    with mp.VideoFileClip(video_path) as video:
        if not video.audio:
            raise ValueError("No audio track found in video file")
        video.audio.write_audiofile(audio_path, logger=None)
    if not os.path.exists(audio_path):
        raise RuntimeError("Audio extraction failed")

def transcribe_audio(audio_path):
    result = model_whisper.transcribe(audio_path)
    return result['text']

def read_text_file(file_path):
    if file_path.lower().endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_path.lower().endswith('.docx'):
        doc = Document(file_path)
        return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        raise ValueError("Unsupported file format. Please use .txt or .docx")

def split_into_sentences(text):
    text = re.sub(r'\s+', ' ', text).strip()
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if s and len(s.strip()) > 0]

def compute_similarity_scores(agenda_text, sentences, model):
    agenda_embedding = model.encode(agenda_text, convert_to_tensor=True)
    sentence_embeddings = model.encode(sentences, convert_to_tensor=True)
    return util.cos_sim(agenda_embedding, sentence_embeddings)[0].cpu().numpy()

def determine_threshold(scores):
    if len(scores) < 2:
        return float(np.percentile(scores, 75)) if scores else 0.0
    X = scores.reshape(-1, 1)
    kmeans = KMeans(n_clusters=2, random_state=0).fit(X)
    return np.mean(kmeans.cluster_centers_)

def create_report(agenda, aligned, non_aligned, threshold):
    alignment_percent = (len(aligned) / (len(aligned) + len(non_aligned))) * 100 if (len(aligned) + len(non_aligned)) > 0 else 0
    report = [
        "Agenda Alignment Analysis Report",
        "=" * 50,
        f"Agenda Topic: {agenda}",
        f"Alignment Percentage: {alignment_percent:.1f}%",
        f"Similarity Threshold: {threshold:.3f}",
        "\nKey Aligned Content:",
        "-" * 30
    ]
    for sent, score in aligned:
        report.append(f"► {sent} (Score: {score:.3f})")
    report.extend([
        "\nNon-Aligned Content:",
        "-" * 30
    ])
    for sent, score in non_aligned:
        report.append(f"◌ {sent} (Score: {score:.3f})")
    return "\n".join(report)

def save_docx_report(content, output_path):
    doc = Document()
    lines = content.split('\n')
    title = lines[0]
    doc.add_heading(title, 0)
    for line in lines[1:]:
        if ':' in line and not line.startswith(('►', '◌')):
            key, value = line.split(':', 1)
            p = doc.add_paragraph()
            p.add_run(key + ':').bold = True
            p.add_run(value)
        elif line.startswith('►'):
            p = doc.add_paragraph(line[2:], style='ListBullet')
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        elif line.startswith('◌'):
            p = doc.add_paragraph(line[2:], style='ListBullet')
        else:
            doc.add_paragraph(line)
    doc.save(output_path)

def process_input(input_type, file_path, agenda, request_id):
    temp_dir = tempfile.gettempdir()
    if input_type == 'video':
        audio_path = os.path.join(temp_dir, f"{request_id}_audio.wav")
        extract_audio_from_video(file_path, audio_path)
        text = transcribe_audio(audio_path)
        os.remove(audio_path)
    elif input_type == 'audio':
        text = transcribe_audio(file_path)
    elif input_type == 'text':
        text = read_text_file(file_path)
    else:
        raise ValueError("Invalid input type")
    
    sentences = split_into_sentences(text)
    if not sentences:
        raise ValueError("No valid sentences found in the input")
    
    scores = compute_similarity_scores(agenda, sentences, model_st)
    threshold = determine_threshold(scores)
    
    aligned = [(s, sc) for s, sc in zip(sentences, scores) if sc >= threshold]
    non_aligned = [(s, sc) for s, sc in zip(sentences, scores) if sc < threshold]
    
    # Calculate alignment percentage
    alignment_percent = (len(aligned) / (len(aligned) + len(non_aligned))) * 100 if (len(aligned) + len(non_aligned)) > 0 else 0
    
    report_content = create_report(agenda, aligned, non_aligned, threshold)
    report_path = os.path.join(temp_dir, f"{request_id}_report.docx")
    save_docx_report(report_content, report_path)
    return report_path, aligned, non_aligned, threshold, alignment_percent

# Store some analysis results for the dashboard
recent_analyses = []

@app.route('/')
def dashboard():
    """Render the main dashboard page"""
    return render_template('dashboard.html', analyses=recent_analyses)

@app.route('/analyze', methods=['GET', 'POST'])
def analyze():
    """Handle the analysis form page"""
    if request.method == 'POST':
        input_type = request.form.get('input_type')
        agenda = request.form.get('agenda', '').strip()
        if 'file' not in request.files:
            return render_template('analyze.html', error="No file uploaded")
        file = request.files['file']
        if file.filename == '':
            return render_template('analyze.html', error="No selected file")
        if not agenda:
            return render_template('analyze.html', error="Please enter an agenda topic")
        
        # Validate file extension
        allowed_extensions = {
            'video': ['.mp4', '.avi', '.mov'],
            'audio': ['.mp3', '.wav', '.m4a'],
            'text': ['.txt', '.docx']
        }
        ext = os.path.splitext(file.filename)[1].lower()
        if input_type not in allowed_extensions or ext not in allowed_extensions[input_type]:
            return render_template('analyze.html', error=f"Invalid file type for {input_type}")
        
        # Save the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            request_id = str(uuid.uuid4())
            report_path, aligned, non_aligned, threshold, alignment_percent = process_input(input_type, file_path, agenda, request_id)
            
            # Store analysis summary for dashboard
            if len(recent_analyses) >= 5:  # Keep only last 5 analyses
                recent_analyses.pop(0)
                
            recent_analyses.append({
                'id': request_id,
                'filename': filename,
                'agenda': agenda,
                'alignment_percent': alignment_percent,
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'type': input_type
            })
            
            # Store report path for download
            report_paths[request_id] = report_path
            
            # Store analysis results
            analysis_results[request_id] = {
                'aligned': aligned,
                'non_aligned': non_aligned,
                'threshold': threshold,
                'alignment_percent': alignment_percent,
                'agenda': agenda,
                'filename': filename
            }
            
            # Redirect to results page
            return redirect(url_for('show_results', analysis_id=request_id))
        except Exception as e:
            return render_template('analyze.html', error=str(e))
        finally:
            # Cleanup uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)
    
    return render_template('analyze.html')

@app.route('/results/<analysis_id>')
def show_results(analysis_id):
    """Show analysis results in browser"""
    # Check if analysis exists
    if analysis_id not in analysis_results:
        return redirect(url_for('dashboard'))
    
    # Get analysis data
    data = analysis_results[analysis_id]
    
    return render_template('results.html', 
                          aligned=data['aligned'],
                          non_aligned=data['non_aligned'],
                          threshold=data['threshold'],
                          alignment_percent=data['alignment_percent'],
                          agenda=data['agenda'],
                          analysis_id=analysis_id)

@app.route('/download-report/<analysis_id>')
def download_report(analysis_id):
    """Download the report as a Word document"""
    if analysis_id in report_paths and os.path.exists(report_paths[analysis_id]):
        return send_file(
            report_paths[analysis_id],
            as_attachment=True,
            download_name="alignment_report.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return redirect(url_for('dashboard'))

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    app.run(debug=True)